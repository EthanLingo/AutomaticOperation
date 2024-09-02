from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH


def add_heading_numbering(input_file, output_file):
    """
    为新生成的Word文档中的所有标题添加正确的编号。

    参数:
    input_file (str): 输入文档的文件路径
    output_file (str): 输出文档的文件路径
    """
    # 打开输入文档
    doc = Document(input_file)

    # 初始化标题编号
    heading_nums = [0] * 9  # 最多支持9级标题

    # 遍历文档中的段落,为所有标题添加编号
    for paragraph in doc.paragraphs:
        for i in range(1, 10):
            if paragraph.style.name == f'Heading {i}':
                heading_nums[i-1] += 1
                # 重置下级标题的编号计数
                for j in range(i, 9):
                    heading_nums[j] = 0
                number_str = '.'.join(map(str, [num for num in heading_nums[:i+1] if num > 0]))
                paragraph.text = f"{number_str}. {paragraph.text}"
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                break

    # 保存输出文档
    doc.save(output_file)


if __name__ == '__main__':
    # 示例用法
    input_file = '你的路径/GraduationThesis/毕业论文之方案二/毕业论文之方案二_0.docx'
    output_file = '你的路径/GraduationThesis/毕业论文之方案二/毕业论文之方案二_1.docx'

    add_heading_numbering(input_file, output_file)
